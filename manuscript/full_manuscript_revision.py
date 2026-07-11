"""
Step 4: Full manuscript revision for npj Microgravity submission.
Rewrites abstract, expands intro, updates results with Fig 8 + supp figs,
expands discussion, adds figure legends, supp fig/table legends, new refs,
updates Data Availability, writes cover letter.

Writes to the WORK scratch dir first (docx is ZIP-based, needs a local path).
"""
# --- portable paths (de-sandboxed; replaces /mnt/results and /workspace) ---
import os as _os
_HERE = _os.path.dirname(_os.path.abspath(__file__))
REPO_ROOT = _os.path.abspath(_os.path.join(_HERE, '..'))
RESULTS = _os.environ.get("ASO_ROOT", REPO_ROOT)          # holds tables/ and figures/
ATLAS   = _os.environ.get("ASO_ATLAS", _os.path.join(REPO_ROOT, "atlas"))  # large intermediates (not shipped)
WORK    = _os.environ.get("ASO_WORK", _os.path.join(REPO_ROOT, "work"))    # scratch outputs
_os.makedirs(WORK, exist_ok=True)
# --- end portable paths ---

from docx import Document
from docx.shared import Pt, Inches
from copy import deepcopy
import os

# Load existing manuscript
doc = Document(WORK + '/manuscript_temp.docx')

def insert_paragraph_after(paragraph, text, style='Normal'):
    """Insert a new paragraph after the given paragraph."""
    new_p = deepcopy(paragraph._element)
    for child in list(new_p):
        if child.tag.endswith('}r'):
            new_p.remove(child)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = doc.styles[style]
    if text:
        new_para.add_run(text)
    return new_para

def find_para_index(text_startswith, start=0):
    for i in range(start, len(doc.paragraphs)):
        if doc.paragraphs[i].text.startswith(text_startswith):
            return i
    return None

def replace_para_text(idx, new_text):
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)

# ============================================================
# 4a. Rewrite Abstract (≤150 words)
# ============================================================
abstract_idx = find_para_index("Spaceflight imposes unique")
new_abstract = (
    "Spaceflight imposes mechanical, oxidative, and gravitational stresses on plants, "
    "but the cell-type-specific signaling architecture underlying plant spaceflight "
    "responses remains poorly characterized. We integrated bulk RNA-seq biomarker "
    "discovery (156 samples, 6 NASA OSDR studies) with single-cell analysis of the "
    "Arabidopsis seed-to-seed atlas (41,314 cells) using scPlantLLM, PlantCellChat, "
    "and ggPlantmap. The LASSO panel (AUC=0.734, 85 genes) revealed stress/ER-proteostasis "
    "enrichment in flight-upregulated genes and flavonoid metabolism in ground-upregulated "
    "genes. PlantCellChat identified Ca2+ as the dominant signaling pathway (10.2-fold "
    "over K+), with guard cells as the primary source. A CBL9-CIPK23-AKT1 crosstalk "
    "circuit connects Ca2+ sensing to K+ uptake, showing strong guard-cell-specific "
    "co-expression (CDPK4-ANNAT1 r=0.56, CBL9-AKT1 r=0.34). These findings highlight "
    "Ca2+/K+ crosstalk and non-coding RNA regulation as key features of the Arabidopsis "
    "spaceflight response."
)
replace_para_text(abstract_idx, new_abstract)

# ============================================================
# 4b. Introduction — add Ca2+/K+ context paragraph
# ============================================================
intro_last_idx = find_para_index("Here we integrate these approaches")
# Update the "Here we integrate" paragraph to include point (6)
new_intro_close = (
    "Here we integrate these approaches to: (1) build a LASSO biomarker panel predicting "
    "spaceflight versus ground conditions from pooled OSDR bulk RNA-seq data, (2) apply "
    "scPlantLLM to the seed-to-seed atlas for zero-shot cell annotation and clustering, "
    "(3) infer cell-cell communication with PlantCellChat, (4) visualize results on plant "
    "anatomy with ggPlantmap, (5) integrate the biomarker panel with the single-cell "
    "signaling architecture, and (6) characterize the Ca2+/K+ crosstalk circuit connecting "
    "Ca2+ signaling to K+ transport via the CBL9-CIPK23-AKT1 cascade. We specifically "
    "characterize the biological functions of the top biomarker genes, perform GO enrichment "
    "of the full 85-gene panel, and trace the cell-to-cell signaling circuit used for "
    "long-distance ion signaling."
)
replace_para_text(intro_last_idx, new_intro_close)

# Insert Ca2+ gravity-sensing context paragraph before "Here we integrate"
intro_context = (
    "Ca2+ is a primary second messenger in plant gravity perception. Microgravity alters "
    "cytosolic Ca2+ oscillations and gradients that guide directional growth, and space "
    "experiments have confirmed that cytosolic Ca2+ increases in response to changes in "
    "gravisensor positioning [14]. The CBL-CIPK calcium signaling cascade is a "
    "well-established downstream effector, and the CBL9-CIPK23-AKT1 pathway specifically "
    "links Ca2+ sensing to K+ uptake, regulating stomatal aperture and root K+ acquisition "
    "[12,13,15]. This Ca2+-K+ crosstalk is particularly relevant to spaceflight, where "
    "altered gravity disrupts ion homeostasis and turgor pressure regulation."
)
insert_paragraph_after(doc.paragraphs[intro_last_idx - 1], intro_context, 'Normal')

# ============================================================
# 4c. Results — update GO enrichment to reference supp figs
# ============================================================
go_results_idx = find_para_index("To characterize the biological processes")
go_text = doc.paragraphs[go_results_idx].text
# Add supp fig reference at end
if "Supplementary Figure" not in go_text:
    go_text += (
        " The directed analysis reveals a clear biological split: spaceflight-upregulated "
        "genes are enriched for stress/heat-shock/ER-proteostasis pathways, while "
        "ground-upregulated genes are enriched for flavonoid/secondary metabolism "
        "(Supplementary Figures S1-S2, Supplementary Tables S2-S3)."
    )
    replace_para_text(go_results_idx, go_text)

# ============================================================
# 4c. Results — update Ca2+/K+ section to reference Figure 8 and Supp Fig S3
# ============================================================
ca2k_idx = find_para_index("To investigate whether the Ca2+-dominant")
ca2k_text = doc.paragraphs[ca2k_idx].text
# Add Figure 8 and Supp Fig references
if "Figure 8" not in ca2k_text:
    # Insert reference to Figure 8 (Sankey) and Supplementary Figure S3 (montage)
    ca2k_text = ca2k_text.replace(
        "Cell-type proxy mapping was therefore used as the more defensible approach.",
        "Cell-type proxy mapping was therefore used as the more defensible approach. "
        "A Sankey diagram tracing signal flow at two levels — cell-type communication "
        "strengths and the molecular CBL9-CIPK23-AKT1 cascade with guard cell expression "
        "values — confirms that guard cells are the dominant Ca2+ source and that the "
        "cascade converges on AKT1-mediated K+ uptake (Figure 8). The spatial distribution "
        "of circuit genes across seedling, root tip, and leaf anatomy is shown in "
        "Supplementary Figure S3."
    )
    replace_para_text(ca2k_idx, ca2k_text)

# ============================================================
# 4d. Discussion — expand Ca2+/K+ paragraph
# ============================================================
disc_ca2k_idx = find_para_index("The Ca2+/K+ crosstalk circuit analysis reveals")
new_disc_ca2k = (
    "The Ca2+/K+ crosstalk circuit analysis reveals a guard cell-centered signaling "
    "module that may serve as the cell-to-cell conduit for long-distance ion signaling "
    "under spaceflight conditions. The CBL9-CIPK23-AKT1 cascade, a canonical "
    "Ca2+-regulated K+ uptake pathway [12,13], provides a mechanistic link between the "
    "Ca2+-dominant CCC network and K+ homeostasis. The strong co-expression of "
    "CDPK4-ANNAT1 and CIPK23-CaM3 specifically in guard cells (r = 0.56 and 0.46, "
    "respectively) suggests that Ca2+ decoding and annexin-mediated Ca2+ buffering "
    "are co-regulated in stomatal lineage cells. Guard cells control stomatal aperture, "
    "which regulates gas exchange and water balance — functions critical in the closed "
    "atmosphere of spacecraft [15]. The 10-fold dominance of Ca2+ over K+ in "
    "intercellular communication strength further supports Ca2+ as the primary "
    "long-range signal, with K+ transport acting as a downstream effector. This "
    "convergence of Ca2+ signaling and ER proteostasis (identified in the GO enrichment "
    "of flight-upregulated genes) suggests a coordinated stress response: Ca2+ "
    "oscillations trigger both CBL-CIPK-mediated ion transport and heat-shock protein "
    "induction to maintain proteostasis under microgravity-induced stress [14]. "
    "However, the absence of organ labels in the seed-to-seed atlas limits spatial "
    "resolution; future studies using organ-specific single-cell datasets will be needed "
    "to validate the tissue-level circuit architecture."
)
replace_para_text(disc_ca2k_idx, new_disc_ca2k)

# ============================================================
# 4e. Methods — update Ca2+/K+ section to mention Sankey and montage
# ============================================================
methods_ca2k_idx = find_para_index("The Ca2+ pathway network was expanded")
methods_ca2k_text = doc.paragraphs[methods_ca2k_idx].text
if "Sankey" not in methods_ca2k_text:
    methods_ca2k_text += (
        " A two-panel Sankey diagram was constructed using Plotly to visualize signal "
        "flow at both the cell-type level (Ca2+ and K+ communication strengths between "
        "5 cell types) and the molecular cascade level (CBL9-CIPK23-AKT1 with guard cell "
        "expression values as flow widths). A supplementary ggPlantmap montage combined "
        "the seedling, root tip, and leaf maps into a single three-panel figure."
    )
    replace_para_text(methods_ca2k_idx, methods_ca2k_text)

# ============================================================
# 4f. Figure Legends — update Figure 5, add Figure 8, add Supp legends
# ============================================================
# Update Figure 5 legend
fig5_legend_idx = find_para_index("a) Data-driven Ca2+ signaling network")
new_fig5_legend = (
    "a) Data-driven Ca2+ signaling network (20 nodes) visualized using the ggpathway "
    "workflow (ggraph + igraph) with stress-majorization layout (graphlayouts::layout_with_stress). "
    "Nodes represent Ca2+ (yellow, ligand), CBL sensors (blue: CBL1, CBL9, CBL2), CIPK kinases "
    "(teal: CIPK23, CIPK1, CIPK9), CaM/CML sensors (purple: CML38, CML42, CML24, CaM7, CaM3), "
    "CDPK kinases (green: CDPK3, CDPK4), downstream targets (orange: CAX3, NRT2.1, SOS1, ANNAT1), "
    "and LASSO biomarker genes (red: HSP22.0 and TRAF-like). Node size indicates mean expression "
    "in the atlas. Edge width indicates communication strength from PlantCellChat. b) Conceptual "
    "schematic of the Ca2+ signaling cascade under microgravity, showing Ca2+ channel activation, "
    "cytosolic Ca2+ elevation, and branching to CBL-CIPK, calmodulin, and CDPK pathways leading "
    "to stress response genes."
)
replace_para_text(fig5_legend_idx, new_fig5_legend)

# Add Figure 8 legend after Figure 7 legend
fig7_legend_idx = find_para_index("a) Expanded Ca2+/K+ pathway network")
if fig7_legend_idx:
    fig7_para = doc.paragraphs[fig7_legend_idx]
    # Find the end of Figure 7 legend (next paragraph that starts with "Figure" or a heading)
    fig7_end = fig7_legend_idx
    for i in range(fig7_legend_idx + 1, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith('Figure') or doc.paragraphs[i].style.name.startswith('Heading'):
            fig7_end = i - 1
            break
    else:
        fig7_end = len(doc.paragraphs) - 1

    fig8_title = insert_paragraph_after(doc.paragraphs[fig7_end], "", 'Normal')
    fig8_title.add_run("Figure 8. Ca2+/K+ signaling circuit Sankey diagram.")

    fig8_text = (
        "a) Cell-type signal flow: Sankey diagram showing Ca2+ (blue) and K+ (green) "
        "communication strengths between five cell types (source on left, target on right). "
        "Flow width is proportional to PlantCellChat communication strength. Guard cells "
        "(orange) are the dominant Ca2+ source; epidermal cells are the primary K+ source. "
        "b) Molecular cascade: Sankey diagram tracing the CBL9-CIPK23-AKT1 pathway from "
        "Ca2+ entry through calcium sensors (CBL1, CBL9, CaM3, CML24) to kinases (CIPK23, "
        "CDPK3, CDPK4) to effectors (AKT1, ANNAT1, KC1) and K+ uptake. Flow widths are "
        "proportional to guard cell expression values. Edge labels show guard cell "
        "co-expression correlations (e.g., CBL9-AKT1 r=0.34, CIPK23-CaM3 r=0.46, "
        "CDPK4-ANNAT1 r=0.56). The orange CIPK23-AKT1 edge highlights the Ca2+-K+ "
        "crosstalk connection."
    )
    insert_paragraph_after(fig8_title, fig8_text, 'Normal')

# Add Supplementary Figure and Table Legends section
# Find the Data Availability heading
data_avail_idx = find_para_index("Data Availability")
if data_avail_idx:
    # Insert supplementary legends section before Data Availability
    supp_heading = insert_paragraph_after(doc.paragraphs[data_avail_idx - 1], "", 'Heading 1')
    supp_heading.add_run("Supplementary Figure Legends")

    supp_figs = [
        ("Supplementary Figure S1. GO enrichment dotplot.",
         "Dot plot showing significant GO and KEGG terms enriched in the 85-gene LASSO panel, "
         "41 flight-up genes, and 44 ground-up genes. Dot size = gene count; color = -log10(p-value). "
         "Background: 2,000 variable features. Correction: BH-FDR (padj < 0.05)."),
        ("Supplementary Figure S2. GO enrichment directed bar plot.",
         "Bar plot showing -log10(p-value) for significant terms, separated by direction "
         "(flight-up vs. ground-up). Flight-up genes are enriched for heat response, hypoxia, "
         "and ER protein processing; ground-up genes are enriched for flavonoid biosynthesis."),
        ("Supplementary Figure S3. Spatial expression of Ca2+/K+ circuit genes across "
         "Arabidopsis anatomy.",
         "Three-panel ggPlantmap montage: a) Seedling map (cotyledon, hypocotyl, root) showing "
         "expression of 6 key circuit genes (AKT1, CBL9, CIPK23, CML24, KC1, CDPK3). "
         "b) Root tip longitudinal section showing Ca2+ cascade gene expression across 12 root "
         "zones. c) Leaf cross-section showing K+ channel distribution across 8 leaf regions. "
         "Cell-type expression was used as a proxy for organ mapping (cotyledon = mesophyll + "
         "guard; root = stele + meristematic)."),
        ("Supplementary Figure S4. Ca2+/K+ co-expression heatmap.",
         "25-gene Pearson correlation matrix across all cells (left) and per cell type (right). "
         "Guard cells (n=89) show the strongest co-expression module: CDPK4-ANNAT1 (r=0.56), "
         "CIPK23-CaM3 (r=0.46), CBL9-AKT1 (r=0.34). Other cell types show weak correlations (r<0.1)."),
        ("Supplementary Figure S5. Ca2+/K+ gene expression per cell type.",
         "Heatmap showing mean expression (left) and percentage expressing (right) for 25 Ca2+/K+ "
         "circuit genes across 5 cell types. KC1 shows highest expression (24.7% expressing); "
         "GORK shows lowest (0.9%)."),
        ("Supplementary Figure S6. Cell-to-cell Ca2+/K+ communication circuit.",
         "Circular network diagram showing Ca2+ (blue) and K+ (green) signaling strength between "
         "5 cell types. Edge width = communication strength. Guard cells are the dominant Ca2+ "
         "source (total outgoing 57.1); epidermal cells are the primary K+ source."),
        ("Supplementary Figure S7. Organ marker diagnostic.",
         "Expression of 15 tissue-specific markers across 5 cell types. Root markers (SCR, PLT1, "
         "WOL) are near-absent (<3%); photosynthetic markers (RBCS, FSD1) are ubiquitous (93-98%); "
         "seed markers (LEC2, 2S3) are near-zero. Marker-based organ inference is not reliable in "
         "this whole-seedling atlas."),
        ("Supplementary Figure S8. Expanded Ca2+/K+ pathway network.",
         "Full 27-node, 67-edge Ca2+/K+ pathway network visualized with ggpathway (stress-majorization "
         "layout). Ca2+ edges in blue, K+ edges in green, crosstalk edge (CBL9-AKT1) in orange. "
         "Node types: CBL sensors, CIPK kinases, CaM/CML sensors, CDPKs, K+ channels, Ca2+ targets, "
         "LASSO biomarker genes."),
    ]

    current = supp_heading
    for title, legend in supp_figs:
        current = insert_paragraph_after(current, title, 'Normal')
        current = insert_paragraph_after(current, legend, 'Normal')

    # Supplementary Table Legends
    tbl_heading = insert_paragraph_after(current, "", 'Heading 1')
    tbl_heading.add_run("Supplementary Table Legends")

    supp_tables = [
        ("Supplementary Table S1. LASSO biomarker panel.",
         "85 stable LASSO features (selection frequency >=50%) with mean coefficient, "
         "standard deviation, and stability status."),
        ("Supplementary Table S2. GO enrichment results (all terms).",
         "Complete g:Profiler output for all 3 queries (85 genes, 41 flight-up, 44 ground-up) "
         "with all tested terms, p-values, and gene lists."),
        ("Supplementary Table S3. GO enrichment significant terms.",
         "14 significant terms (BH-FDR padj < 0.05) across all 3 queries with -log10(p) and "
         "gene lists."),
        ("Supplementary Table S4. Ca2+/K+ CCC circuit.",
         "25 cell-pair communication strengths for Ca2+ and K+ pathways, with organ proxies "
         "and dominance flags."),
        ("Supplementary Table S5. Ca2+/K+ signaling circuit summary.",
         "Circuit-level summary: source/target cell types, dominant signal, Ca2+/K+ strength, "
         "ratio, top source/target genes, and strongest co-expression pair."),
        ("Supplementary Table S6. Ca2+/K+ circuit node expression.",
         "23-gene expression table: mean expression and percentage expressing across 5 cell "
         "types, with top co-expression partner and correlation."),
        ("Supplementary Table S7. Ca2+/K+ co-expression matrix.",
         "25x25 Pearson correlation matrix for Ca2+/K+ circuit genes (global, across all cells)."),
        ("Supplementary Table S8. Organ marker expression.",
         "15 tissue-specific markers tested across 5 cell types with mean expression and "
         "percentage expressing."),
    ]

    current = tbl_heading
    for title, legend in supp_tables:
        current = insert_paragraph_after(current, title, 'Normal')
        current = insert_paragraph_after(current, legend, 'Normal')

# ============================================================
# 4g. References — add refs 14-15
# ============================================================
ref13_idx = find_para_index("13. Xu, J.")
if ref13_idx:
    ref13_para = doc.paragraphs[ref13_idx]

    ref14 = insert_paragraph_after(ref13_para, "", 'Normal')
    ref14.add_run(
        "14. Takahashi, K., Takahashi, H., Furuichi, T., Toyota, M., Furutani-Seiki, M., "
        "Kobayashi, T., Watanabe-Takano, H. & Shinohara, M. Gravity sensing in plant and "
        "animal cells. npj Microgravity 7, 2 (2021)."
    )

    ref15 = insert_paragraph_after(ref14, "", 'Normal')
    ref15.add_run(
        "15. Cheong, Y. H., Pandey, G. K., Grant, J. J., Batistic, O., Li, L., Kim, B.-G., "
        "Lee, S.-C., Kudla, J. & Luan, S. Two calcineurin B-like calcium sensors, interacting "
        "with protein kinase CIPK23, regulate leaf transpiration and root potassium uptake "
        "in Arabidopsis. Plant J. 52, 473-484 (2007)."
    )

# ============================================================
# 4h. Data Availability — update with repo/Zenodo
# ============================================================
data_avail_text_idx = find_para_index("OSDR data are available")
new_data_avail = (
    "OSDR data are available from the NASA Open Science Data Repository "
    "(visualization.osdr.nasa.gov/biodata/api/). The seed-to-seed atlas is available from "
    "GEO (accession GSE226097). scPlantLLM is available from GitHub "
    "(github.com/compbioNJU/scPlantLLM). PlantCellChat is available from GitHub "
    "(github.com/mrliuw/PlantCellChat). ggPlantmap is available from GitHub "
    "(github.com/leonardojo/ggPlantmap). All analysis code, generated figures, data tables, "
    "and supplementary materials are available in the companion repository "
    "(github.com/rbarker/arabidopsis-spaceflight-omics) and will be archived on Zenodo "
    "with a DOI upon acceptance. Detailed methods including the scPlantLLM weight conversion "
    "code, all analysis parameters, and software versions are provided in Supplementary Methods."
)
replace_para_text(data_avail_text_idx, new_data_avail)

# ============================================================
# Save manuscript
# ============================================================
doc.save(WORK + '/manuscript_revised.docx')
print("Manuscript revised and saved to " + WORK + "/manuscript_revised.docx")

# Verify
doc2 = Document(WORK + '/manuscript_revised.docx')
print(f"Total paragraphs: {len(doc2.paragraphs)}")
headings = [(i, p.style.name, p.text[:60]) for i, p in enumerate(doc2.paragraphs) if p.style.name.startswith('Heading')]
for i, style, text in headings:
    print(f"  {i}: [{style}] {text}")

# Word count
words = sum(len(p.text.split()) for p in doc2.paragraphs)
body_words = 0
for p in doc2.paragraphs:
    if p.style.name.startswith('Heading') and ('References' in p.text or 'Figure Legends' in p.text or 'Data Availability' in p.text or 'Acknowledgements' in p.text or 'Author Contributions' in p.text or 'Competing' in p.text or 'Supplementary' in p.text):
        break
    body_words += len(p.text.split())
print(f"\nTotal words: {words}")
print(f"Body words (through Methods): ~{body_words}")

# Abstract word count
abstract_p = doc2.paragraphs[find_para_index("Spaceflight imposes mechanical")]
print(f"Abstract words: {len(abstract_p.text.split())}")
