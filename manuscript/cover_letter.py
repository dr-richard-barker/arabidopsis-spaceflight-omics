"""
Step 4 (continued): Generate cover letter for npj Microgravity submission.
"""
# --- portable paths (de-sandboxed; replaces /mnt/results and /workspace) ---
import os as _os
_HERE = _os.path.dirname(_os.path.abspath(__file__))
REPO_ROOT = _os.path.abspath(_os.path.join(_HERE, '..'))
RESULTS = _os.environ.get("ASO_ROOT", REPO_ROOT)          # holds tables/ and figures/
ATLAS   = _os.environ.get("ASO_ATLAS", _os.path.join(REPO_ROOT, "atlas"))  # large intermediates (not shipped)
WORK    = _os.environ.get("ASO_WORK", _os.path.join(REPO_ROOT, "work"))    # scratch outputs
_os.makedirs(WORK, exist_ok=True)
# --- end portable paths ---

from docx import Document
from docx.shared import Pt, Inches

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles['Normal']
style.font.name = 'Liberation Sans'
style.font.size = Pt(11)

doc.add_paragraph('July 8, 2026')
doc.add_paragraph('')
doc.add_paragraph('Editorial Team')
doc.add_paragraph('npj Microgravity')
doc.add_paragraph('Nature Portfolio')
doc.add_paragraph('')
doc.add_paragraph('Dear Editor,')
doc.add_paragraph('')
p = doc.add_paragraph(
    "We submit our manuscript entitled \"LASSO biomarker panel and single-cell atlas "
    "integration reveals Ca2+-dominated cell-cell communication in Arabidopsis "
    "spaceflight response\" for consideration as a Research Article in npj Microgravity."
)
p = doc.add_paragraph(
    "This study integrates bulk spaceflight transcriptomics from the NASA Open Science "
    "Data Repository (156 samples across 6 studies) with single-cell analysis of the "
    "Arabidopsis seed-to-seed atlas (41,314 cells) to connect biomarker-level findings "
    "with cell-type-resolved signaling architecture. We believe three findings are "
    "particularly significant for the npj Microgravity readership:"
)
p = doc.add_paragraph(
    "1. A LASSO biomarker panel (AUC=0.734, 85 genes) that generalizes across ISS "
    "seedling studies, with GO enrichment revealing a biological split: flight-upregulated "
    "genes are enriched for stress/ER-proteostasis pathways, while ground-upregulated "
    "genes are enriched for flavonoid metabolism."
)
p = doc.add_paragraph(
    "2. Ca2+ is the dominant cell-cell communication pathway (10.2-fold over K+), "
    "consistent with its role as a primary gravity-perception second messenger. "
    "Guard cells are the dominant Ca2+ signaling source."
)
p = doc.add_paragraph(
    "3. A novel Ca2+/K+ crosstalk circuit — the CBL9-CIPK23-AKT1 cascade — that "
    "mechanistically connects Ca2+ sensing to K+ uptake, with strong guard-cell-specific "
    "co-expression (CDPK4-ANNAT1 r=0.56, CBL9-AKT1 r=0.34). This circuit may serve as "
    "the cell-to-cell conduit for long-distance ion signaling under spaceflight conditions, "
    "with implications for stomatal regulation and turgor pressure in closed spacecraft "
    "atmospheres."
)
p = doc.add_paragraph(
    "All analysis code, data tables, and figures are available in a companion GitHub "
    "repository (github.com/rbarker/arabidopsis-spaceflight-omics) with full "
    "reproducibility documentation, and will be archived on Zenodo with a DOI upon "
    "acceptance. The manuscript has not been published elsewhere and is not under "
    "consideration by another journal."
)
p = doc.add_paragraph(
    "We believe this work is well-suited for npj Microgravity as it advances "
    "understanding of the molecular signaling architecture underlying plant responses "
    "to the spaceflight environment, with direct relevance to space agriculture and "
    "life support systems."
)
doc.add_paragraph('')
doc.add_paragraph('Sincerely,')
doc.add_paragraph('')
doc.add_paragraph('Richard Barker')
doc.add_paragraph('GeneLab Plant Analysis Working Group')
doc.add_paragraph('NASA Open Science Data Repository')
doc.add_paragraph('rbarker@nasa.gov')

doc.save(WORK + '/cover_letter.docx')
print("Cover letter saved to " + WORK + "/cover_letter.docx")
